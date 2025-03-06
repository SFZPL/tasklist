import os
import io
import ast
import xmlrpc.client
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from datetime import datetime, time, timedelta, date
from typing import Tuple, Optional, List, Dict, Any

# =========================================
# Load Environment Variables and Set Up Odoo
# =========================================
load_dotenv()
ODOO_URL = st.secrets["odoo"]["ODOO_URL"]
ODOO_DB = st.secrets["odoo"]["ODOO_DB"]
ODOO_USERNAME = st.secrets["odoo"]["ODOO_USERNAME"]
ODOO_PASSWORD = st.secrets["odoo"]["ODOO_PASSWORD"]

@st.cache_resource(show_spinner=False)
def get_odoo_connection() -> Tuple[Optional[int], Optional[xmlrpc.client.ServerProxy]]:
    try:
        common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
        uid = common.authenticate(ODOO_DB, ODOO_USERNAME, ODOO_PASSWORD, {})
        if not uid:
            st.error("Failed to authenticate with Odoo. Check credentials and DB name.")
            return None, None
        models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
        return uid, models
    except Exception as e:
        st.error(f"Error connecting to Odoo: {e}")
        return None, None

# =========================================
# Shared Helper Functions
# =========================================
def parse_domain(domain_string: str) -> List:
    try:
        return ast.literal_eval(domain_string)
    except Exception:
        return []

def combine_domains_or(domains_list: List[List]) -> List:
    if not domains_list:
        return []
    combined = domains_list[0]
    for d in domains_list[1:]:
        combined = ['|'] + combined + d
    return combined

def get_field_name(many2one_val: Any) -> str:
    if isinstance(many2one_val, list) and len(many2one_val) > 1:
        return many2one_val[1]
    return ""

def format_datetime_range(start_dt_str: str, end_dt_str: str) -> str:
    """Converts times (assumed GMT) to GMT+3 and returns the formatted range."""
    try:
        if start_dt_str:
            dt_start = datetime.strptime(start_dt_str, "%Y-%m-%d %H:%M:%S") + timedelta(hours=3)
            start_fmt = dt_start.strftime("%Y-%m-%d %H:%M:%S")
        else:
            start_fmt = ""
        if end_dt_str:
            dt_end = datetime.strptime(end_dt_str, "%Y-%m-%d %H:%M:%S") + timedelta(hours=3)
            end_fmt = dt_end.strftime("%Y-%m-%d %H:%M:%S")
        else:
            end_fmt = ""
    except Exception:
        # Fallback to original strings if parsing fails.
        start_fmt = start_dt_str
        end_fmt = end_dt_str
    if not end_fmt:
        return start_fmt
    return f"{start_fmt} -> {end_fmt}"

# =========================================
# (A) MORNING TASK LIST (from planning.slot)
# =========================================
def get_planning_favorites(models, uid) -> List[Dict[str, Any]]:
    try:
        domain = [('model_id', '=', 'planning.slot')]
        fields_to_read = ['name', 'domain']
        return models.execute_kw(
            ODOO_DB, uid, ODOO_PASSWORD,
            'ir.filters', 'search_read',
            [domain],
            {'fields': fields_to_read}
        )
    except Exception as e:
        st.error(f"Error retrieving favorites: {e}")
        return []

def get_tasks(models, uid, final_domain: List) -> List[Dict[str, Any]]:
    fields_to_read = [
        'resource_id',
        'role_id',
        'x_studio_parent_task',
        'x_studio_sub_task_1',
        'start_datetime',
        'end_datetime',
        'state',
        'allocated_hours'
    ]
    try:
        return models.execute_kw(
            ODOO_DB, uid, ODOO_PASSWORD,
            'planning.slot', 'search_read',
            [final_domain],
            {'fields': fields_to_read}
        )
    except Exception as e:
        st.error(f"Error retrieving tasks: {e}")
        return []

def fetch_subtask_details(models, uid: int, subtask_ids: List[int]) -> Dict[int, Dict[str, Any]]:
    if not subtask_ids:
        return {}
    subtask_ids = list(set(subtask_ids))
    fields_to_read = [
        'x_studio_service_category_1',
        'x_studio_total_no_of_design_units_sc1'
    ]
    try:
        records = models.execute_kw(
            ODOO_DB, uid, ODOO_PASSWORD,
            'project.task', 'read',
            [subtask_ids],
            {'fields': fields_to_read}
        )
        return {r['id']: r for r in records}
    except Exception as e:
        st.error(f"Error retrieving sub-task details: {e}")
        return {}

def build_morning_text(task: Dict[str, Any], subtask_map: Dict[int, Dict[str, Any]]) -> str:
    lines = []
    # Role
    role_name = get_field_name(task.get('role_id'))
    if role_name:
        lines.append(f"Role: {role_name}")
    # Parent Task
    parent_name = get_field_name(task.get('x_studio_parent_task'))
    if not parent_name:
        lines.append("Parent Task: Missing Parent Task")
    else:
        lines.append(f"Parent Task: {parent_name}")
        # Sub Task
        sub_val = task.get('x_studio_sub_task_1')
        sub_name = get_field_name(sub_val)
        if sub_name:
            lines.append(f"Sub Task: {sub_name}")
            sub_id = sub_val[0] if isinstance(sub_val, list) and len(sub_val) > 0 else None
            if sub_id and sub_id in subtask_map:
                sub_rec = subtask_map[sub_id]
                # Format service category: if it's a many2one, extract the second element
                sc = sub_rec.get('x_studio_service_category_1', '')
                if isinstance(sc, list):
                    sc = get_field_name(sc)
                units = sub_rec.get('x_studio_total_no_of_design_units_sc1', '')
                if sc:
                    lines.append(f"Service Category: {sc}")
                if units:
                    lines.append(f"No. of Units: {units}")
    # Date Range (convert to GMT+3)
    start_dt_str = task.get('start_datetime') or ""
    end_dt_str = task.get('end_datetime') or ""
    drange = format_datetime_range(start_dt_str, end_dt_str)
    if drange:
        lines.append(f"Date Range: {drange}")
    return "\n".join(lines)

def create_morning_table(doc: Document, 
                         tasks_by_designer: Dict[str, List[Dict[str, Any]]],
                         subtask_map: Dict[int, Dict[str, Any]]) -> bytes:
    """
    Creates a 2-column table with one row per designer.
    The second column concatenates all tasks (sorted by start time) for that designer.
    """
    doc.add_heading("Morning Task List", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Designer"
    hdr_cells[1].text = "Task Details"
    
    # For each designer, sort tasks by start_datetime (earlier first)
    for designer, tasks in tasks_by_designer.items():
        try:
            tasks.sort(key=lambda t: datetime.strptime(t.get('start_datetime', "1970-01-01 00:00:00"), "%Y-%m-%d %H:%M:%S"))
        except Exception:
            pass  # In case of parsing error, leave unsorted.
        row_cells = table.add_row().cells
        row_cells[0].text = designer
        task_texts = [build_morning_text(task, subtask_map) for task in tasks]
        row_cells[1].text = "\n\n".join(task_texts)
    
    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()

# =========================================
# (B) RECAP from x_recaps (Unchanged)
# =========================================
def get_recaps(models, uid, date_domain: List) -> List[Dict[str, Any]]:
    fields_to_read = [
        'create_uid',
        'x_studio_shift',
        'x_studio_recap_cat',
        'x_studio_designer_summary',
        'create_date',
        'x_studio_parent_task',
        'x_studio_subtask'
    ]
    try:
        recs = models.execute_kw(
            ODOO_DB, uid, ODOO_PASSWORD,
            'x_recaps', 'search_read',
            [date_domain],
            {'fields': fields_to_read}
        )
        return recs
    except Exception as e:
        st.error(f"Error retrieving recaps from x_recaps: {e}")
        return []

def build_recap_notes_text(rec: Dict[str, Any]) -> str:
    lines = []
    pt_val = rec.get('x_studio_parent_task')
    if isinstance(pt_val, list) and len(pt_val) > 1:
        lines.append(f"Parent Task: {pt_val[1]}")
    st_val = rec.get('x_studio_subtask')
    if st_val:
        lines.append(f"Sub Task: {st_val}")
    shift = rec.get('x_studio_shift', '')
    if shift:
        lines.append(f"Shift: {shift}")
    rc = rec.get('x_studio_recap_cat', '')
    if rc:
        lines.append(f"Recap Category: {rc}")
    cmt = rec.get('x_studio_designer_summary', '')
    if cmt:
        lines.append(f"Comment: {cmt}")
    dt = rec.get('create_date', '')
    if dt:
        lines.append(f"Date & Time: {dt}")
    return "\n".join(lines)

def create_recap_notes_table(doc: Document, recs_by_designer: Dict[str, List[Dict[str, Any]]]) -> bytes:
    doc.add_heading("Recap", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Designer"
    hdr_cells[1].text = "Recap Details"
    for designer, recs in recs_by_designer.items():
        for r in recs:
            row_cells = table.add_row().cells
            row_cells[0].text = designer
            row_cells[1].text = build_recap_notes_text(r)
    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()

# =========================================
# (C) Helper: Get Employees Matching Favorites (for Recaps)
# =========================================
def get_employees_from_favorites(models, uid, fav_domains: List[List]) -> set:
    """For each favorite domain, fetch hr.employee records and return a set of employee names."""
    emp_names = set()
    for domain in fav_domains:
        try:
            emps = models.execute_kw(
                ODOO_DB, uid, ODOO_PASSWORD,
                'hr.employee', 'search_read',
                [domain],
                {'fields': ['name']}
            )
            for emp in emps:
                name = emp.get('name')
                if name:
                    emp_names.add(name)
        except Exception as e:
            st.error(f"Error retrieving employees for domain {domain}: {e}")
    return emp_names

# =========================================
# MAIN
# =========================================
def main():
    st.set_page_config(
        page_title="Planner Tasks & Recaps - Combined",
        page_icon=":clipboard:",
        layout="centered"
    )
    st.title("Task Extraction Tool")
    
    uid, models = get_odoo_connection()
    if not uid:
        st.stop()

    report_type = st.selectbox("Select Report Type", ["Morning Task List", "Recap"])
    
    st.subheader("Select Date Range")
    start_date = st.date_input("Start Date", value=datetime.today())
    end_date = st.date_input("End Date", value=datetime.today())
    # For planning.slot, filter on start_datetime; for recaps, on create_date
    start_dt_str = datetime.combine(start_date, time(0, 0, 0)).strftime("%Y-%m-%d %H:%M:%S")
    end_dt_str = datetime.combine(end_date, time(23, 59, 59)).strftime("%Y-%m-%d %H:%M:%S")
    
    if report_type == "Morning Task List":
        st.write("Optional: Select Favorites for planning.slot.")
        all_favs = get_planning_favorites(models, uid)
        fav_names = [f["name"] for f in all_favs] if all_favs else []
        selected_favs = st.multiselect("Select Favorites (optional)", fav_names)
        
        if st.button("Fetch & Generate Morning Tasks"):
            fav_domains = []
            for fav in selected_favs:
                rec = next((x for x in all_favs if x["name"] == fav), None)
                if rec:
                    d = parse_domain(rec.get("domain", "[]"))
                    if d:
                        fav_domains.append(d)
            combined_fav_domain = combine_domains_or(fav_domains)
            date_domain = [
                ('start_datetime', '>=', start_dt_str),
                ('start_datetime', '<=', end_dt_str)
            ]
            final_domain = (['&'] + date_domain + combined_fav_domain) if combined_fav_domain else date_domain
            tasks = get_tasks(models, uid, final_domain)
            if tasks:
                st.success(f"Fetched {len(tasks)} tasks from planning.slot!")
                # Group tasks by designer and sort them by start_datetime (earlier first)
                tasks_by_designer = {}
                for t in tasks:
                    res = t.get('resource_id')
                    designer = res[1] if (res and isinstance(res, list) and len(res) > 1) else "Unassigned"
                    tasks_by_designer.setdefault(designer, []).append(t)
                for designer in tasks_by_designer:
                    try:
                        tasks_by_designer[designer].sort(
                            key=lambda t: datetime.strptime(t.get('start_datetime', "1970-01-01 00:00:00"), "%Y-%m-%d %H:%M:%S")
                        )
                    except Exception:
                        pass
                subtask_ids = []
                for t in tasks:
                    sub_val = t.get('x_studio_sub_task_1')
                    if isinstance(sub_val, list) and len(sub_val) > 0:
                        subtask_ids.append(sub_val[0])
                subtask_map = fetch_subtask_details(models, uid, subtask_ids)
                doc = Document()
                doc_bytes = create_morning_table(doc, tasks_by_designer, subtask_map)
                st.download_button(
                    label="Download Morning Tasks",
                    data=doc_bytes,
                    file_name="Morning_Task_List.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("No tasks found with the selected filters (Morning).")
    else:
        # Recap (x_recaps) branch remains unchanged
        st.write("Optional: Select Favorites for Recaps (filter by creator).")
        all_favs = get_planning_favorites(models, uid)
        fav_names = [f["name"] for f in all_favs] if all_favs else []
        selected_favs = st.multiselect("Select Favorites (optional)", fav_names)
        fav_domains = []
        for fav in selected_favs:
            rec = next((x for x in all_favs if x["name"] == fav), None)
            if rec:
                d = parse_domain(rec.get("domain", "[]"))
                if d:
                    fav_domains.append(d)
        allowed_emp_names = get_employees_from_favorites(models, uid, fav_domains) if fav_domains else None
        
        if st.button("Fetch & Generate Recap from x_recaps"):
            recap_domain = [
                ('create_date', '>=', start_dt_str),
                ('create_date', '<=', end_dt_str)
            ]
            recs = get_recaps(models, uid, recap_domain)
            if recs:
                recs_by_designer = {}
                for r in recs:
                    c_uid = r.get('create_uid')
                    designer_name = c_uid[1] if (isinstance(c_uid, list) and len(c_uid) > 1) else "Unassigned"
                    if allowed_emp_names is not None and designer_name not in allowed_emp_names:
                        continue
                    recs_by_designer.setdefault(designer_name, []).append(r)
                if recs_by_designer:
                    doc = Document()
                    doc_bytes = create_recap_notes_table(doc, recs_by_designer)
                    st.download_button(
                        label="Download Recap Report",
                        data=doc_bytes,
                        file_name="Recap.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No recaps found matching the selected favorites.")
            else:
                st.error("No recaps found in x_recaps for that date range.")

if __name__ == "__main__":
    main()
